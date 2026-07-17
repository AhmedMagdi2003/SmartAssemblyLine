from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
README_PATH = OUTPUT_DIR / "README.md"
DOCX_PATH = OUTPUT_DIR / "SmartAssemblyLine_Database_Guide.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5A, 0x6A, 0x7A)
BLACK = RGBColor(0x11, 0x11, 0x11)
LIGHT_FILL = "F7F9FC"
NOTE_FILL = "EEF4FB"
TABLE_HEADER_FILL = "E8EEF5"
BORDER = "B7C4D6"


def dxa_from_inches(inches: float) -> int:
    return int(round(inches * 1440))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size=8) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_in_inches, indent_dxa=120) -> None:
    widths_dxa = [dxa_from_inches(width) for width in widths_in_inches]
    total_width = sum(widths_dxa)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_inline_runs(paragraph, text: str, font="Calibri", size=11, color=BLACK) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=max(size - 0.5, 9), color=BLACK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, name=font, size=size, color=color)


def set_paragraph_style(paragraph, before=0, after=6, line=1.25, left=0.0, first_line=0.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.left_indent = Inches(left)
    paragraph.paragraph_format.first_line_indent = Inches(first_line)


def add_paragraph_block(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_style(paragraph, before=0, after=6, line=1.25)
    add_inline_runs(paragraph, text)


def add_bullet_list(doc, items) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_style(paragraph, before=0, after=4, line=1.25)
        add_inline_runs(paragraph, item)


def add_numbered_list(doc, items) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph_style(paragraph, before=0, after=4, line=1.25)
        add_inline_runs(paragraph, item)


def add_note_box(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5], indent_dxa=120)
    set_table_borders(table, color="9FB6D4", size=10)
    cell = table.cell(0, 0)
    shade_cell(cell, NOTE_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_style(paragraph, before=0, after=0, line=1.15)
    add_inline_runs(paragraph, text, size=10.5, color=DARK_BLUE)
    doc.add_paragraph()


def add_code_block(doc, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5], indent_dxa=120)
    set_table_borders(table, color="D6DEE8", size=8)
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    cell.text = ""
    for line in code.strip("\n").splitlines():
        paragraph = cell.add_paragraph()
        set_paragraph_style(paragraph, before=0, after=1, line=1.15)
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9.5, color=BLACK)
    if cell.paragraphs and not cell.paragraphs[0].text:
        cell._tc.remove(cell.paragraphs[0]._p)
    doc.add_paragraph()


def add_table_block(doc, title: str, headers, rows, widths, indent_dxa=120) -> None:
    title_paragraph = doc.add_paragraph()
    set_paragraph_style(title_paragraph, before=4, after=4, line=1.15)
    run = title_paragraph.add_run(title)
    set_run_font(run, name="Calibri", size=10.5, color=MUTED, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent_dxa=indent_dxa)
    set_table_borders(table)

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        shade_cell(cell, TABLE_HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_style(paragraph, before=0, after=0, line=1.1)
        run = paragraph.add_run(header)
        set_run_font(run, name="Calibri", size=10.5, color=DARK_BLUE, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_style(paragraph, before=0, after=0, line=1.15)
            add_inline_runs(paragraph, str(value), size=10.5)

    doc.add_paragraph()


@dataclass(frozen=True)
class TableBlock:
    title: str
    headers: list[str]
    rows: list[list[str]]
    widths: list[float]
    indent_dxa: int = 120


def paragraph(text: str) -> dict:
    return {"type": "paragraph", "text": text}


def bullets(items: list[str]) -> dict:
    return {"type": "bullets", "items": items}


def numbers(items: list[str]) -> dict:
    return {"type": "numbers", "items": items}


def code(text: str) -> dict:
    return {"type": "code", "text": text}


def note(text: str) -> dict:
    return {"type": "note", "text": text}


def table_block(table: TableBlock) -> dict:
    return {"type": "table", "table": table}


GUIDE = {
    "title": "Smart Assembly Line Database Guide",
    "subtitle": "How to inspect, query, reset, rebuild, and extract data from the project database",
    "metadata": [
        ("Project", "SmartAssemblyLine"),
        ("Audience", "Developers and operators working on the local factory stack"),
        ("Source basis", "Current workspace code and docs reviewed on 2026-07-08"),
    ],
    "sections": [
        {
            "level": 1,
            "title": "1. Database Mental Model",
            "blocks": [
                note(
                    "Important repo note: this checkout contains `deployment/docker-compose.local.yml`. "
                    "Use `deployment/docker-compose.local.yml` for the real local database and MQTT stack."
                ),
                paragraph(
                    "At runtime, this project is `PostgreSQL`-only. The vision pipeline publishes carton completion events over `MQTT`, "
                    "the logger writes those events into Postgres and mirrors them into CSV, and the FastAPI dashboard reads historical data back from the database."
                ),
                bullets(
                    [
                        "The runtime source of truth is Postgres, not the CSV files.",
                        "The only runtime table in the current schema is `box_events`.",
                        "The dashboard mixes live updates over MQTT/WebSocket with database-backed history and analytics endpoints.",
                        "Duplicate events are prevented by a unique `uuid` on each carton event.",
                    ]
                ),
                table_block(
                    TableBlock(
                        title="Runtime Facts",
                        headers=["Setting", "Value"],
                        rows=[
                            ["Compose file", "`deployment/docker-compose.local.yml`"],
                            ["Database container", "`smart-assembly-db`"],
                            ["MQTT container", "`smart-assembly-mqtt`"],
                            ["Database image", "`pgvector/pgvector:0.8.2-pg17`"],
                            ["Database URL", "`postgresql://smartassembly:smartassembly@localhost:5433/smart_assembly`"],
                            ["Database name", "`smart_assembly`"],
                            ["User / password", "`smartassembly / smartassembly`"],
                            ["Compose DB volume key", "`smart_assembly_db_data` (Docker will project-prefix the real volume name)"],
                            ["Backup trail", "`data/logs/shift_<ShiftName>_<YYYY-MM-DD>.csv`"],
                        ],
                        widths=[1.875, 4.625],
                    )
                ),
            ],
        },
        {
            "level": 1,
            "title": "2. File Map You Will Actually Use",
            "blocks": [
                paragraph(
                    "These are the main files that matter when you need to understand, query, or rebuild the database layer."
                ),
                table_block(
                    TableBlock(
                        title="Database File Map",
                        headers=["Path", "Why it matters"],
                        rows=[
                            ["`src/db/settings.py`", "Loads `.env` and resolves `DATABASE_URL`."],
                            ["`src/db/session.py`", "Builds the SQLAlchemy `engine`, `SessionLocal`, and `get_db()` helper."],
                            ["`src/db/models.py`", "Defines the `BoxEvent` ORM model and its columns."],
                            ["`src/db/repositories.py`", "Contains the read/write functions used by the logger and dashboard."],
                            ["`src/db/bootstrap.py`", "Runs `Base.metadata.create_all()` for local bootstrapping."],
                            ["`scripts/init_db.py`", "Small entry point that calls `create_database()`."],
                            ["`alembic/env.py`", "Points Alembic at the same runtime database URL and metadata."],
                            ["`alembic/versions/20260412_0001_create_box_events.py`", "Current tracked schema migration."],
                            ["`scripts/fetch_db.py`", "CLI tool for extracting rows, summaries, and counts without writing SQL."],
                            ["`src/utils/logger.py`", "Subscribes to MQTT, inserts rows into Postgres, and mirrors them to CSV."],
                            ["`src/dashboard/main.py`", "Reads history and KPIs from repository functions and exposes API endpoints."],
                            ["`deployment/docker-compose.local.yml`", "Starts the local Postgres and Mosquitto containers."],
                        ],
                        widths=[2.15, 4.35],
                    )
                ),
            ],
        },
        {
            "level": 1,
            "title": "3. What The Database Stores",
            "blocks": [
                paragraph(
                    "The current schema stores completed carton events in one table named `box_events`. "
                    "The `uuid` is the important business key, and it encodes the operational shift date and shift name."
                ),
                table_block(
                    TableBlock(
                        title="`box_events` Columns",
                        headers=["Column", "Type", "Meaning"],
                        rows=[
                            ["`id`", "integer", "Primary key used for internal ordering."],
                            ["`uuid`", "string(128)", "Unique event key such as `BOX-20260404-Morning_Shift-0001`."],
                            ["`yolo_session_id`", "integer", "Tracker ID at the moment the carton was completed."],
                            ["`timestamp_iso`", "string(64)", "Production timestamp stored as text and used in ordering."],
                            ["`shift`", "string(64)", "Logical shift name such as `Morning_Shift`."],
                            ["`shift_count`", "integer", "Running carton count inside the shift window."],
                            ["`transit_time_sec`", "float", "Measured carton travel time."],
                            ["`orientation_deg`", "float", "Measured orientation angle from the vision step."],
                            ["`status`", "string(32)", "Current code writes `COMPLETED`."],
                            ["`created_at`", "timestamp with time zone", "Database-side insertion timestamp from `now()`."],
                        ],
                        widths=[1.35, 1.45, 3.70],
                    )
                ),
                bullets(
                    [
                        "The runtime app is Postgres-only even though some unit tests use SQLite as a lightweight test harness.",
                        "The effective `shift_date` used by filters and KPIs is derived from the `uuid` prefix, not from `created_at`.",
                        "For night-shift data that crosses midnight, the operational date still comes from the encoded `uuid` date. This is intentional and is also covered by tests.",
                        "The logger writes CSV only after a successful new database insert. If the same `uuid` is seen again, the duplicate is skipped.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "4. Start, Stop, And Health-Check The Database",
            "blocks": [
                {"type": "heading", "level": 2, "title": "4.1 Start only the database infrastructure"},
                paragraph("This is the minimum you need if you want Postgres and MQTT available but you are not launching the full app stack."),
                code(
                    """docker compose -f deployment/docker-compose.local.yml up -d
python -m alembic upgrade head
python -c "from src.db.session import DATABASE_URL; print(DATABASE_URL)" """
                ),
                {"type": "heading", "level": 2, "title": "4.2 Start the local stack the repo already provides"},
                paragraph("The PowerShell launcher starts Docker, waits for health, applies Alembic migrations, then launches the logger and dashboard."),
                code(
                    r"""powershell -ExecutionPolicy Bypass -File scripts/start_local_stack.ps1
powershell -ExecutionPolicy Bypass -File scripts/start_local_stack.ps1 -NoPipeline"""
                ),
                {"type": "heading", "level": 2, "title": "4.3 Stop the stack without deleting data"},
                paragraph("The provided stop script shuts down background services and runs `docker compose down`, but it does not erase the Postgres volume."),
                code(
                    r"""powershell -ExecutionPolicy Bypass -File scripts/stop_local_stack.ps1
docker compose -f deployment/docker-compose.local.yml down"""
                ),
                {"type": "heading", "level": 2, "title": "4.4 Quick health checks"},
                code(
                    """docker ps --filter "name=smart-assembly"
docker inspect -f "{{if .State.Health}}{{.State.Health.Status}}{{else}}{{.State.Status}}{{end}}" smart-assembly-db
python scripts/fetch_db.py count
curl http://127.0.0.1:8000/api/health"""
                ),
                bullets(
                    [
                        "If `api/health` says `database=connected`, the dashboard can reach the same database URL the app is using.",
                        "If `scripts/fetch_db.py count` works but the dashboard is empty, the dashboard process may be running with a different environment or before migrations were applied.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "5. Extract Data Without Writing SQL",
            "blocks": [
                paragraph(
                    "The easiest extraction tool already in the repo is `scripts/fetch_db.py`. "
                    "It uses the same SQLAlchemy session layer as the main app, so it is the safest first step for day-to-day inspection."
                ),
                {"type": "heading", "level": 2, "title": "5.1 Read recent or specific events"},
                code(
                    """python scripts/fetch_db.py events --limit 20
python scripts/fetch_db.py events --latest
python scripts/fetch_db.py events --shift Morning_Shift --shift-date 2026-04-19 --limit 50
python scripts/fetch_db.py events --uuid BOX-20260419-Morning_Shift-0007
python scripts/fetch_db.py events --limit 5 --json"""
                ),
                {"type": "heading", "level": 2, "title": "5.2 Read summaries and counts"},
                code(
                    """python scripts/fetch_db.py shifts --limit 10
python scripts/fetch_db.py count
python scripts/fetch_db.py count --shift Morning_Shift --shift-date 2026-04-19"""
                ),
                bullets(
                    [
                        "Use `--json` when you want structured output that is easy to redirect into a file or another script.",
                        "The `--shift-date` filter matches the encoded operational date in the UUID prefix, which is why it works correctly for overnight shifts.",
                    ]
                ),
                {"type": "heading", "level": 2, "title": "5.3 Read through the dashboard API"},
                code(
                    """curl "http://127.0.0.1:8000/api/events?limit=20"
curl "http://127.0.0.1:8000/api/events/latest"
curl "http://127.0.0.1:8000/api/kpis/current"
curl "http://127.0.0.1:8000/api/stats/shifts?limit=10"
curl "http://127.0.0.1:8000/api/charts/overview?limit=100"
curl "http://127.0.0.1:8000/api/events?current_shift_only=true&limit=100"
curl "http://127.0.0.1:8000/api/charts/overview?shift=Morning_Shift&shift_date=2026-04-19&limit=100" """
                ),
                bullets(
                    [
                        "The API is useful when you want the same shape the browser dashboard consumes.",
                        "In `/api/events`, rows are fetched newest-first in the repository layer and then reversed before the API response, so the API payload is oldest-to-newest inside the selected result set.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "6. Extract Data With Direct SQL",
            "blocks": [
                paragraph(
                    "When `fetch_db.py` is not enough, open `psql` inside the Postgres container and run direct SQL against `box_events`."
                ),
                code("""docker exec -it smart-assembly-db psql -U smartassembly -d smart_assembly"""),
                {"type": "heading", "level": 2, "title": "6.1 Latest rows"},
                code(
                    """SELECT id, uuid, shift, shift_count, transit_time_sec, orientation_deg, timestamp_iso
FROM box_events
ORDER BY id DESC
LIMIT 20;"""
                ),
                {"type": "heading", "level": 2, "title": "6.2 One exact shift window"},
                code(
                    """SELECT id, uuid, shift_count, transit_time_sec, orientation_deg, timestamp_iso
FROM box_events
WHERE shift = 'Morning_Shift'
  AND uuid LIKE 'BOX-20260419-Morning_Shift-%'
ORDER BY id;"""
                ),
                {"type": "heading", "level": 2, "title": "6.3 Summary by operational shift window"},
                code(
                    """SELECT split_part(uuid, '-', 2) AS shift_date_raw,
       shift,
       COUNT(*) AS events,
       MAX(shift_count) AS shift_volume,
       ROUND(AVG(transit_time_sec)::numeric, 2) AS avg_transit_sec
FROM box_events
GROUP BY split_part(uuid, '-', 2), shift
ORDER BY shift_date_raw DESC, shift;"""
                ),
                {"type": "heading", "level": 2, "title": "6.4 Find slow or misaligned cartons"},
                code(
                    """SELECT uuid, shift, shift_count, transit_time_sec, orientation_deg, timestamp_iso
FROM box_events
WHERE transit_time_sec > 2.0
   OR ABS(orientation_deg) > 15
ORDER BY timestamp_iso DESC;"""
                ),
                bullets(
                    [
                        "Use `MAX(shift_count)` when you want final volume for a shift window, because `shift_count` is cumulative inside that window.",
                        "If you need file exports, the safest repo-native path is usually `python scripts/fetch_db.py ... --json` and then redirect the output in your shell.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "7. Extract Data From Python",
            "blocks": [
                paragraph(
                    "Use the repository helpers when you want to write one-off analysis scripts without re-implementing filters and serializers."
                ),
                {"type": "heading", "level": 2, "title": "7.1 Repository-level reads"},
                code(
                    """from src.db.repositories import get_current_kpis, get_shift_summary, list_recent_box_events

print(get_current_kpis())
print(get_shift_summary(limit=3))

for row in list_recent_box_events(limit=5, shift='Morning_Shift', shift_date='2026-04-19'):
    print(row['uuid'], row['transit_time_sec'])"""
                ),
                {"type": "heading", "level": 2, "title": "7.2 Raw SQLAlchemy session reads"},
                code(
                    """from src.db.models import BoxEvent
from src.db.session import SessionLocal

session = SessionLocal()
try:
    rows = (
        session.query(BoxEvent)
        .order_by(BoxEvent.id.desc())
        .limit(10)
        .all()
    )
    for row in rows:
        print(row.id, row.uuid, row.transit_time_sec)
finally:
    session.close()"""
                ),
                bullets(
                    [
                        "Use repository functions when you want the same semantics as the dashboard and logger.",
                        "Use raw sessions when you need custom SQLAlchemy queries the repository layer does not already expose.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "8. Reset The Database Safely",
            "blocks": [
                note(
                    "Before any reset, stop the writer processes first: the vision pipeline, the logger, and the dashboard. "
                    "If they stay running, they can immediately reconnect and start repopulating the database while you are trying to clear it."
                ),
                {"type": "heading", "level": 2, "title": "8.1 Optional: take a plain SQL backup first"},
                code("""docker exec smart-assembly-db pg_dump -U smartassembly smart_assembly > smart_assembly_backup.sql"""),
                {"type": "heading", "level": 2, "title": "8.2 Reset only the data, keep the schema"},
                paragraph("This is the fastest safe reset when the table shape is fine and you only want an empty history."),
                code(
                    """docker exec -i smart-assembly-db psql -U smartassembly -d smart_assembly -c "TRUNCATE TABLE box_events RESTART IDENTITY;" """
                ),
                {"type": "heading", "level": 2, "title": "8.3 Reset the schema through Alembic"},
                paragraph(
                    "Use this when you want to drop tracked tables and recreate them from the migration history without destroying the Docker container itself."
                ),
                code(
                    """python -m alembic downgrade base
python -m alembic upgrade head"""
                ),
                {"type": "heading", "level": 2, "title": "8.4 Wipe the whole local database volume and rebuild from zero"},
                paragraph(
                    "This is the most complete local reset. It removes the Docker-managed Postgres volume, starts a fresh container, and reapplies migrations."
                ),
                code(
                    """docker compose -f deployment/docker-compose.local.yml down -v
docker compose -f deployment/docker-compose.local.yml up -d
python -m alembic upgrade head"""
                ),
                bullets(
                    [
                        "`docker compose ... down -v` removes the Postgres volume and the Mosquitto volumes declared in the same compose file.",
                        "`scripts/stop_local_stack.ps1` does not erase database data because it calls plain `docker compose down`.",
                        "`python scripts/init_db.py` is useful for local bootstrapping, but for a real repo-consistent rebuild you should prefer `python -m alembic upgrade head`.",
                    ]
                ),
                {"type": "heading", "level": 2, "title": "8.5 Recommended reset choices"},
                table_block(
                    TableBlock(
                        title="Which reset should you choose?",
                        headers=["Situation", "Best reset"],
                        rows=[
                            ["You only want to clear history", "`TRUNCATE TABLE box_events RESTART IDENTITY;`"],
                            ["You changed migrations and want a clean reapply", "`python -m alembic downgrade base` then `python -m alembic upgrade head`"],
                            ["You want a truly fresh local DB container", "`docker compose -f deployment/docker-compose.local.yml down -v` then `up -d` and `upgrade head`"],
                        ],
                        widths=[2.65, 3.85],
                    )
                ),
            ],
        },
        {
            "level": 1,
            "title": "9. Change The Schema And Keep The Project Healthy",
            "blocks": [
                paragraph(
                    "When you add new columns or tables, change the ORM model first, generate a migration, inspect it, then upgrade the database."
                ),
                numbers(
                    [
                        "Edit the SQLAlchemy model under `src/db/models.py`.",
                        "Generate a migration with `python -m alembic revision --autogenerate -m \"describe your change\"`.",
                        "Inspect the generated file under `alembic/versions/` before running it.",
                        "Apply the change with `python -m alembic upgrade head`.",
                        "Smoke-test the database with `python scripts/fetch_db.py count` and one or two dashboard endpoints.",
                    ]
                ),
                code(
                    """python -m alembic revision --autogenerate -m "add new column"
python -m alembic upgrade head
python scripts/fetch_db.py count
curl http://127.0.0.1:8000/api/health"""
                ),
                bullets(
                    [
                        "The current tracked migration is `alembic/versions/20260412_0001_create_box_events.py`.",
                        "Alembic uses the same database URL resolver as the app because `alembic/env.py` calls `get_database_url()`.",
                        "If you skip migration review, it is easy to create a diff that technically runs but does not match the intended schema.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "10. Troubleshooting And Useful Real-World Notes",
            "blocks": [
                bullets(
                    [
                        "If you get `DATABASE_URL is not set`, either create `.env` in the project root or export `DATABASE_URL` manually before starting the process.",
                        "If the logger falls back to CSV-only behavior, the database layer was unavailable when `src/utils/logger.py` started.",
                        "If the dashboard health endpoint says `database=error`, test the same DB URL with `python scripts/fetch_db.py count` from the same environment.",
                        "If you see fewer rows than expected, remember that duplicate UUIDs are ignored on insert by design.",
                        "If you want a clean shutdown without losing data, use the stop script or `docker compose down` without `-v`.",
                        "If you want a real wipe, `docker compose down -v` is the command that matters. Stopping containers alone is not enough.",
                        "If you are comparing output from the API and direct SQL, remember the API may reverse event order for the frontend experience.",
                    ]
                ),
            ],
        },
        {
            "level": 1,
            "title": "11. Quick Command Checklist",
            "blocks": [
                code(
                    """# Start local infrastructure
docker compose -f deployment/docker-compose.local.yml up -d
python -m alembic upgrade head

# Check the database quickly
python scripts/fetch_db.py count
curl http://127.0.0.1:8000/api/health

# Read events
python scripts/fetch_db.py events --limit 20
python scripts/fetch_db.py events --limit 20 --json

# Empty the table but keep the schema
docker exec -i smart-assembly-db psql -U smartassembly -d smart_assembly -c "TRUNCATE TABLE box_events RESTART IDENTITY;"

# Full local wipe and rebuild
docker compose -f deployment/docker-compose.local.yml down -v
docker compose -f deployment/docker-compose.local.yml up -d
python -m alembic upgrade head"""
                ),
            ],
        },
    ],
}


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_style(title, before=0, after=4, line=1.0)
    run = title.add_run(GUIDE["title"])
    set_run_font(run, name="Calibri", size=24, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_style(subtitle, before=0, after=12, line=1.1)
    add_inline_runs(subtitle, GUIDE["subtitle"], size=12, color=MUTED)

    metadata_rows = [[label, value] for label, value in GUIDE["metadata"]]
    add_table_block(
        doc,
        title="Guide Metadata",
        headers=["Field", "Value"],
        rows=metadata_rows,
        widths=[1.55, 4.95],
        indent_dxa=120,
    )


def add_heading(doc: Document, title: str, level: int) -> None:
    doc.add_paragraph(title, style=f"Heading {level}")


def build_docx(path: Path) -> None:
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_style(footer_paragraph, before=0, after=0, line=1.0)
    add_inline_runs(footer_paragraph, "Smart Assembly Line Database Guide", size=9, color=MUTED)

    add_title_block(doc)

    for section_data in GUIDE["sections"]:
        add_heading(doc, section_data["title"], section_data["level"])
        for block in section_data["blocks"]:
            block_type = block["type"]
            if block_type == "paragraph":
                add_paragraph_block(doc, block["text"])
            elif block_type == "bullets":
                add_bullet_list(doc, block["items"])
            elif block_type == "numbers":
                add_numbered_list(doc, block["items"])
            elif block_type == "code":
                add_code_block(doc, block["text"])
            elif block_type == "note":
                add_note_box(doc, block["text"])
            elif block_type == "table":
                table = block["table"]
                add_table_block(doc, table.title, table.headers, table.rows, table.widths, indent_dxa=table.indent_dxa)
            elif block_type == "heading":
                add_heading(doc, block["title"], block["level"])
            else:
                raise ValueError(f"Unknown block type: {block_type}")

    doc.save(path)


def markdown_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    divider = "| " + " | ".join(["---"] * len(headers)) + " |"
    lines = ["| " + " | ".join(headers) + " |", divider]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def build_markdown(path: Path) -> None:
    lines: list[str] = []
    lines.append(f"# {GUIDE['title']}")
    lines.append("")
    lines.append(GUIDE["subtitle"])
    lines.append("")

    lines.extend(markdown_table(["Field", "Value"], [[label, value] for label, value in GUIDE["metadata"]]))
    lines.append("")

    for section_data in GUIDE["sections"]:
        lines.append("#" * section_data["level"] + " " + section_data["title"])
        lines.append("")
        for block in section_data["blocks"]:
            block_type = block["type"]
            if block_type == "paragraph":
                lines.append(block["text"])
                lines.append("")
            elif block_type == "bullets":
                for item in block["items"]:
                    lines.append(f"- {item}")
                lines.append("")
            elif block_type == "numbers":
                for index, item in enumerate(block["items"], start=1):
                    lines.append(f"{index}. {item}")
                lines.append("")
            elif block_type == "code":
                lines.append("```text")
                lines.extend(block["text"].strip("\n").splitlines())
                lines.append("```")
                lines.append("")
            elif block_type == "note":
                lines.append("> " + block["text"])
                lines.append("")
            elif block_type == "table":
                table = block["table"]
                lines.append(f"**{table.title}**")
                lines.append("")
                lines.extend(markdown_table(table.headers, table.rows))
                lines.append("")
            elif block_type == "heading":
                lines.append("#" * block["level"] + " " + block["title"])
                lines.append("")
            else:
                raise ValueError(f"Unknown block type: {block_type}")

    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_markdown(README_PATH)
    build_docx(DOCX_PATH)
    print(f"Wrote {README_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
