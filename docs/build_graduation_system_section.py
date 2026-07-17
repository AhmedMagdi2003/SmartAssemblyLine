from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "SmartAssemblyLine_Graduation_System_Section.docx"


COLOR_HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
COLOR_HEADING_DARK = RGBColor(0x1F, 0x4D, 0x78)
COLOR_INK = RGBColor(0x0B, 0x25, 0x45)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)


TITLE = "Smart Assembly Line: System Design, Rationale, and Engineering Challenges"
SUBTITLE = (
    "Graduation Project Book Section - professional and easy-to-read summary "
    "of the implemented system"
)


SECTIONS = [
    {
        "heading": "1. System Overview",
        "paragraphs": [
            (
                "The Smart Assembly Line is an industrial computer vision and "
                "monitoring system built to observe cartons on a production "
                "conveyor, measure important operational data, and present the "
                "results in real time. The system combines live image capture, "
                "object detection, object tracking, event generation, message "
                "streaming, database persistence, and a browser dashboard in one "
                "connected workflow."
            ),
            (
                "In practical terms, the system detects each carton, follows it "
                "while it moves on the belt, estimates its orientation angle, "
                "measures how long it spends in the monitored area, counts it "
                "inside the active work shift, and then publishes one structured "
                "completion event for that carton. That event is saved to "
                "PostgreSQL, mirrored to CSV as a backup, and displayed live on "
                "the dashboard for operators and supervisors."
            ),
            (
                "The implemented software stack is based on Python, OpenCV, "
                "Ultralytics YOLOv8, BotSort tracking, MQTT, FastAPI, Plotly, "
                "SQLAlchemy, Alembic, PostgreSQL, and Docker Compose. The result "
                "is a modular edge analytics system that is practical for factory "
                "use and still flexible enough for future expansion."
            ),
        ],
    },
    {
        "heading": "2. Why We Built the System",
        "paragraphs": [
            (
                "The main reason for building this project was the need for a "
                "more reliable and more informative way to monitor carton flow on "
                "the assembly line. Manual counting and visual observation are "
                "slow, tiring, and vulnerable to human error, especially during "
                "long shifts or when production speed changes. They also do not "
                "provide structured historical data that can be used later for "
                "analysis or process improvement."
            ),
            (
                "We wanted a system that could count products automatically, but "
                "we also wanted more than a simple counter. In a real production "
                "environment, the team also needs to know whether cartons are "
                "well aligned, whether movement through the line is stable, how "
                "many boxes have been completed in each shift, and whether the "
                "production trend is improving or drifting. For that reason, the "
                "project was designed as an analytics system, not only as a "
                "detection demo."
            ),
            (
                "Our design goals were clear: detect cartons accurately, count "
                "each carton once and only once, measure orientation and transit "
                "time, preserve data after refresh or restart, allow live "
                "operator monitoring, and keep the architecture modular so that "
                "each layer can be improved without rewriting the whole system. "
                "Another important goal was to keep the heavy vision processing "
                "near the factory floor while making the final data easy to move "
                "to dashboards or remote services."
            ),
        ],
    },
    {
        "heading": "3. End-to-End System Workflow",
        "subsections": [
            {
                "heading": "3.1 Image Acquisition",
                "paragraphs": [
                    (
                        "The current runtime is prepared to receive live frames "
                        "from a Raspberry Pi camera sender over the local network. "
                        "During development, recorded conveyor video was also used "
                        "for repeatable testing and algorithm tuning. This staged "
                        "approach allowed us to stabilize the core logic before "
                        "depending fully on live hardware."
                    )
                ],
            },
            {
                "heading": "3.2 Detection and Tracking",
                "paragraphs": [
                    (
                        "Each incoming frame is processed by a YOLOv8 model to "
                        "detect cartons, and BotSort is used to assign persistent "
                        "tracking IDs across consecutive frames. Persistent IDs are "
                        "important because they allow the system to treat a moving "
                        "carton as one physical object instead of many unrelated "
                        "detections."
                    )
                ],
            },
            {
                "heading": "3.3 Filtering and Triggering",
                "paragraphs": [
                    (
                        "After detection, the tracker applies a polygon region of "
                        "interest so only the conveyor area is considered. It also "
                        "checks the box area and the tracked lifespan to reject "
                        "small or short-lived false positives. A carton is counted "
                        "only when its center crosses the finish line and only if "
                        "it has already been tracked for the minimum required time."
                    )
                ],
            },
            {
                "heading": "3.4 Analytics and Payload Generation",
                "paragraphs": [
                    (
                        "When a valid carton crosses the finish line, the system "
                        "calculates its orientation, determines the active shift, "
                        "updates the shift count, and generates a structured event "
                        "payload. The payload includes a unique UUID, the tracker "
                        "session ID, the timestamp, the shift name, the shift "
                        "date, the shift count, the transit time, the orientation "
                        "angle, and the final status."
                    )
                ],
            },
            {
                "heading": "3.5 Messaging and Persistence",
                "paragraphs": [
                    (
                        "The completed payload is published to the MQTT topic "
                        "`factory/assembly/boxes`. A background logger subscribes "
                        "to this topic, saves new events into PostgreSQL, and then "
                        "writes the same event into a daily shift-specific CSV "
                        "file. In this way, the database becomes the main source "
                        "of truth while the CSV files remain a practical backup "
                        "and export path."
                    )
                ],
            },
            {
                "heading": "3.6 Visualization and Monitoring",
                "paragraphs": [
                    (
                        "The FastAPI dashboard serves the operator interface, "
                        "restores historical data from PostgreSQL through API "
                        "endpoints, and receives live MQTT-driven updates through "
                        "a WebSocket bridge. This gives the user both live status "
                        "and durable history in the same interface."
                    )
                ],
            },
        ],
    },
    {
        "heading": "4. System Architecture",
        "paragraphs": [
            (
                "The implemented architecture can be described as a modular "
                "event-driven edge analytics architecture. It is modular because "
                "vision, messaging, persistence, and visualization are separated "
                "into different components. It is event-driven because the system "
                "does not stream every frame into the database. Instead, it emits "
                "one business event when a carton is actually completed."
            ),
            (
                "This architecture contains six main layers. The first layer is "
                "the acquisition layer, where a camera or Raspberry Pi sender "
                "produces image frames. The second layer is the vision layer, "
                "where YOLOv8 and BotSort detect and track cartons. The third "
                "layer is the analytics layer, where shift logic, timing, "
                "counting, and orientation measurements are produced. The fourth "
                "layer is the communication layer, where MQTT decouples the edge "
                "pipeline from downstream services. The fifth layer is the "
                "persistence layer, where PostgreSQL and CSV storage preserve the "
                "result. The sixth layer is the presentation layer, where FastAPI "
                "and Plotly provide the dashboard."
            ),
        ],
        "subsections": [
            {
                "heading": "4.1 Edge Vision Layer",
                "paragraphs": [
                    (
                        "The edge vision layer is responsible for real-time frame "
                        "processing. It is placed close to the camera source so "
                        "that heavy image processing happens locally instead of "
                        "sending raw video to a remote server. This choice reduces "
                        "bandwidth consumption and keeps response time low."
                    )
                ],
            },
            {
                "heading": "4.2 Analytics and Event Layer",
                "paragraphs": [
                    (
                        "The analytics layer converts raw tracking information "
                        "into business-level production events. This layer is "
                        "important because factories do not need millions of raw "
                        "frame records. They need meaningful outputs such as box "
                        "count, orientation, shift volume, and transit time."
                    )
                ],
            },
            {
                "heading": "4.3 Communication Layer",
                "paragraphs": [
                    (
                        "MQTT was selected as the communication layer because it "
                        "is lightweight, simple, and effective for machine-to-"
                        "machine messaging. It allows the tracker, logger, and "
                        "dashboard pipeline to remain loosely coupled. If one "
                        "consumer changes, the producer does not need major "
                        "modification."
                    )
                ],
            },
            {
                "heading": "4.4 Persistence Layer",
                "paragraphs": [
                    (
                        "PostgreSQL is the system's persistent storage layer. It "
                        "stores all completed events in a structured form and "
                        "supports later querying, analytics, and dashboard "
                        "bootstrap. SQLAlchemy and Alembic provide a clear schema "
                        "management path, while the unique UUID constraint helps "
                        "protect the data from duplicates."
                    )
                ],
            },
            {
                "heading": "4.5 Visualization Layer",
                "paragraphs": [
                    (
                        "The visualization layer is built with FastAPI, WebSocket "
                        "broadcasting, and a Plotly-based web interface. It shows "
                        "current shift status, production volume, average transit "
                        "time, orientation drift, and event history. The dashboard "
                        "also supports both current-shift monitoring and all-"
                        "history review."
                    )
                ],
            },
            {
                "heading": "4.6 Deployment and Operations Layer",
                "paragraphs": [
                    (
                        "The system includes local Docker Compose services for "
                        "PostgreSQL and Mosquitto, PowerShell and WSL launcher "
                        "scripts for practical operation, and documented options "
                        "for cloud or tunnel-based remote access. This makes the "
                        "solution easier to run in the lab today and easier to "
                        "expand later."
                    )
                ],
            },
        ],
    },
    {
        "heading": "5. Why This Architecture Was Chosen",
        "paragraphs": [
            (
                "The first reason for choosing this architecture was separation of "
                "concerns. Detection, tracking, analytics, messaging, storage, and "
                "visualization are different technical problems. By separating "
                "them, we made the system easier to debug, test, and improve."
            ),
            (
                "The second reason was efficiency. Instead of sending raw video to "
                "the database or dashboard, the system creates one compact event "
                "per completed carton. This reduces storage overhead and network "
                "traffic while still keeping the most useful operational data."
            ),
            (
                "The third reason was reliability. If the browser refreshes, the "
                "dashboard can restore history from PostgreSQL. If the same event "
                "is accidentally received twice, the database UUID constraint "
                "prevents duplication. If operators need a backup outside the "
                "database, CSV files still exist. These choices make the system "
                "more practical in a real industrial setting."
            ),
            (
                "The fourth reason was scalability. Because MQTT decouples the "
                "components, the vision node can remain local while the dashboard "
                "or database can move later to another machine or to a cloud "
                "environment. The documented hybrid deployment options already "
                "reflect this design direction."
            ),
            (
                "The fifth reason was maintainability. The project includes unit "
                "tests for tracking, analytics, logging, repository logic, MQTT "
                "configuration, and dashboard behavior. A modular architecture "
                "supports this testing approach much better than one large script."
            ),
        ],
    },
    {
        "heading": "6. Main Technical Challenges and How We Solved Them",
        "subsections": [
            {
                "heading": "6.1 Challenge: Background Noise and Irrelevant Detections",
                "paragraphs": [
                    (
                        "In factory footage, the camera sees more than the moving "
                        "cartons. It can also capture floor texture, belt edges, "
                        "support structures, shadows, and other irrelevant areas. "
                        "If the model processed everything equally, the count would "
                        "be less stable."
                    ),
                    (
                        "We solved this by introducing a polygon region of "
                        "interest that isolates the belt area. The project also "
                        "includes a calibration tool that helps define the polygon "
                        "from the real scene. As a result, the tracker focuses on "
                        "the conveyor instead of the whole frame."
                    ),
                ],
            },
            {
                "heading": "6.2 Challenge: False Positives and Short-Lived Detections",
                "paragraphs": [
                    (
                        "Small visual artifacts or unstable detections can appear "
                        "for only a few frames. If every detection is accepted, "
                        "the system can generate incorrect counts."
                    ),
                    (
                        "We addressed this with two filters. First, detections "
                        "must exceed a minimum box area. Second, a carton must "
                        "remain tracked for a minimum lifespan before it is "
                        "eligible for counting. Together, these filters reduce "
                        "noise and keep the event stream cleaner."
                    ),
                ],
            },
            {
                "heading": "6.3 Challenge: Counting the Same Carton More Than Once",
                "paragraphs": [
                    (
                        "A single carton stays visible for many frames, so a naive "
                        "frame-based counter would count it many times. This is "
                        "one of the most important problems in line monitoring."
                    ),
                    (
                        "Our solution combines persistent tracking IDs, a finish "
                        "line trigger, and a one-shot exit registry. Each carton "
                        "is counted only when it crosses the finish line for the "
                        "first time, and the tracker records that the event has "
                        "already been emitted. This makes the counting logic "
                        "repeatable and easy to reason about."
                    ),
                ],
            },
            {
                "heading": "6.4 Challenge: Estimating Orientation Under Real Lighting Conditions",
                "paragraphs": [
                    (
                        "Orientation estimation is harder than simple detection "
                        "because lighting changes, reflections, and partial edges "
                        "can affect the visible carton shape. A single thresholding "
                        "method is often not enough."
                    ),
                    (
                        "We solved this by using a layered OpenCV approach. The "
                        "algorithm enhances the grayscale crop with CLAHE, tries "
                        "edge-based contour extraction first, then falls back to "
                        "multiple binary and inverse threshold variants using Otsu "
                        "and adaptive thresholding. The best plausible contour is "
                        "selected, and the final angle is computed using "
                        "minAreaRect. This makes the orientation logic more stable "
                        "across lighting variation."
                    ),
                ],
            },
            {
                "heading": "6.5 Challenge: Low-Latency Live Streaming From the Camera Source",
                "paragraphs": [
                    (
                        "In a live production scenario, old frames are not useful. "
                        "If buffering becomes large, the operator sees delayed "
                        "events and the dashboard loses real-time value."
                    ),
                    (
                        "To reduce latency, the system uses a custom raw socket "
                        "receiver instead of a simple default video capture path. "
                        "The receiver disables Nagle's algorithm, enlarges the "
                        "receive buffer, uses non-blocking reads, drains the "
                        "available TCP buffer, extracts the latest complete JPEG "
                        "frame, and discards older bytes. This keeps the system "
                        "closer to live behavior."
                    ),
                ],
            },
            {
                "heading": "6.6 Challenge: Losing Count After Restart or Power Failure",
                "paragraphs": [
                    (
                        "A factory system must survive restarts without resetting "
                        "the operational count in the middle of a shift. If the "
                        "software starts again from zero, the production records "
                        "become inconsistent."
                    ),
                    (
                        "We solved this by synchronizing the analytics layer with "
                        "the persisted database state. At shift startup, the "
                        "system loads the last stored shift count for the active "
                        "shift window and continues from there. This approach makes "
                        "same-shift recovery possible after restart."
                    ),
                ],
            },
            {
                "heading": "6.7 Challenge: Night Shift Crossing Midnight",
                "paragraphs": [
                    (
                        "A normal calendar day changes at midnight, but a night "
                        "shift usually continues as one operational period. If the "
                        "system uses only wall-clock date rules, part of one shift "
                        "can be split incorrectly into two different days."
                    ),
                    (
                        "Our solution was to keep an operational shift date and "
                        "encode it directly inside the event UUID. For example, a "
                        "box completed after midnight can still belong to the "
                        "previous night's shift date. The repository layer and the "
                        "logger both use this logic, which preserves continuity for "
                        "overnight production."
                    ),
                ],
            },
            {
                "heading": "6.8 Challenge: Preventing Duplicate Records in Storage",
                "paragraphs": [
                    (
                        "Messaging systems and long-running processes can sometimes "
                        "retry or replay data, especially after restarts or "
                        "network issues. Without protection, duplicates can enter "
                        "the database and break analytics."
                    ),
                    (
                        "We addressed this on two levels. At the application "
                        "level, the logger recognizes duplicate insert attempts and "
                        "skips repeated payloads. At the database level, the "
                        "UUID column is unique, so PostgreSQL becomes the final "
                        "guard against duplication."
                    ),
                ],
            },
            {
                "heading": "6.9 Challenge: Keeping the Dashboard Useful After Refresh",
                "paragraphs": [
                    (
                        "A live WebSocket alone is not enough for a monitoring "
                        "dashboard, because refreshing the page would erase the "
                        "context that the operator had already seen."
                    ),
                    (
                        "We solved this by giving the dashboard a two-stage data "
                        "strategy. On page load, it first requests persisted "
                        "history and KPI data from PostgreSQL through REST "
                        "endpoints. After that, it continues receiving live events "
                        "through a WebSocket path that is fed by MQTT. This "
                        "combination provides both continuity and liveness."
                    ),
                ],
            },
            {
                "heading": "6.10 Challenge: Building a System That Works Locally Today and Can Expand Tomorrow",
                "paragraphs": [
                    (
                        "Industrial student projects often work in the lab but "
                        "become difficult to deploy or extend later. We wanted to "
                        "avoid a design that depends on one hard-coded machine."
                    ),
                    (
                        "For that reason, the project uses environment-based MQTT "
                        "settings, Dockerized infrastructure for local services, "
                        "launch scripts for Windows and WSL, and documented paths "
                        "for cloud or tunnel-based access. The current result is a "
                        "local-first architecture with a clear upgrade path for "
                        "remote monitoring and future distributed deployment."
                    ),
                ],
            },
        ],
    },
    {
        "heading": "7. Strengths of the Final System",
        "paragraphs": [
            (
                "One major strength of the final system is that it transforms raw "
                "video into meaningful production events rather than just drawing "
                "boxes on a screen. This makes it more useful for real operations, "
                "because the output can support reporting, process review, and "
                "decision making."
            ),
            (
                "Another strength is the balance between live monitoring and data "
                "durability. Operators can watch events in real time, while "
                "managers can still return later to historical shift data. This "
                "balance is achieved through the combined use of MQTT, "
                "PostgreSQL, and the FastAPI dashboard."
            ),
            (
                "A third strength is modularity. The project is organized into "
                "clear components for communication, tracking, orientation, "
                "analytics, persistence, and dashboard behavior. This structure "
                "makes the codebase easier to maintain and easier to improve in "
                "future work."
            ),
            (
                "A fourth strength is practical validation. The repository "
                "includes tests for tracking logic, analytics behavior, logger "
                "persistence order, repository filtering, MQTT configuration, "
                "orientation recovery, and dashboard API behavior. These tests do "
                "not replace real factory validation, but they greatly improve "
                "development confidence."
            ),
        ],
    },
    {
        "heading": "8. Current Limitations and Future Improvements",
        "paragraphs": [
            (
                "The current system is focused on monitoring and analytics, not on "
                "closed-loop machine control. It can observe the production line "
                "and report useful metrics, but it does not yet drive actuators or "
                "perform automatic correction on the physical process."
            ),
            (
                "Although the project now supports live camera streaming, local "
                "operation, and documented cloud-ready deployment paths, some "
                "future work is still important. Full end-to-end production "
                "validation with the final hardware path should continue. Security "
                "hardening for MQTT and cloud services can be improved further. "
                "Historical analytics can also be expanded with richer date-range "
                "aggregation, stronger observability, and more formal operational "
                "monitoring."
            ),
            (
                "Another useful future improvement would be stronger integration "
                "testing across the whole path from vision event to MQTT message "
                "to database row to dashboard view. This would complement the "
                "current unit tests and make release confidence even stronger."
            ),
        ],
    },
    {
        "heading": "9. Conclusion",
        "paragraphs": [
            (
                "The Smart Assembly Line project demonstrates how computer vision "
                "can be combined with event messaging, database persistence, and "
                "web monitoring to solve a real industrial problem. Instead of "
                "building an isolated detection script, we built a system that "
                "produces operational value: one-time carton events, shift-aware "
                "counting, orientation analytics, historical storage, and live "
                "visibility."
            ),
            (
                "The most important design decision was to treat the project as a "
                "modular production analytics pipeline. This decision influenced "
                "every technical choice, from the use of persistent tracking IDs, "
                "to the finish-line trigger, to MQTT decoupling, to PostgreSQL "
                "history, to dashboard bootstrap after refresh. Because of these "
                "choices, the final system is easier to trust, easier to extend, "
                "and more suitable for real deployment than a simple proof of "
                "concept."
            ),
            (
                "In summary, the project was built to make production monitoring "
                "more accurate, more informative, and more scalable. The final "
                "implementation shows a complete engineering story: identifying "
                "the operational problem, selecting a suitable architecture, "
                "solving technical challenges one by one, and delivering a system "
                "that can continue growing after the graduation project itself."
            ),
        ],
    },
]


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def ensure_style(document: Document, style_name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    styles = document.styles
    if style_name in styles:
        return styles[style_name]
    return styles.add_style(style_name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title_style = document.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = COLOR_INK
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(6)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_style = ensure_style(document, "Subtitle")
    subtitle_style.font.name = "Calibri"
    subtitle_style.font.size = Pt(12)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = COLOR_MUTED
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(16)
    subtitle_style.paragraph_format.line_spacing = 1.15
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = COLOR_HEADING_BLUE
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.line_spacing = 1.0

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = COLOR_HEADING_BLUE
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing = 1.0

    heading3 = document.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = COLOR_HEADING_DARK
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)
    heading3.paragraph_format.line_spacing = 1.0

    body_lead = ensure_style(document, "Body Lead")
    body_lead.font.name = "Calibri"
    body_lead.font.size = Pt(11)
    body_lead.font.bold = True
    body_lead.font.color.rgb = COLOR_INK
    body_lead.paragraph_format.space_before = Pt(6)
    body_lead.paragraph_format.space_after = Pt(4)
    body_lead.paragraph_format.line_spacing = 1.15


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Smart Assembly Line Graduation Project")
    set_run_font(run, size=9, color=COLOR_MUTED)


def add_title_block(document: Document) -> None:
    document.add_paragraph(TITLE, style="Title")
    document.add_paragraph(SUBTITLE, style="Subtitle")

    intro = document.add_paragraph(style="Body Lead")
    intro_run = intro.add_run(
        "This section explains what we built, why we built it, how the system "
        "works, which technical challenges appeared during development, and how "
        "those challenges were solved in the final implementation."
    )
    set_run_font(intro_run, size=11, color=COLOR_INK, bold=True)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=RGBColor(0, 0, 0))


def add_section(document: Document, section_data: dict) -> None:
    document.add_paragraph(section_data["heading"], style="Heading 1")
    for paragraph in section_data.get("paragraphs", []):
        add_body_paragraph(document, paragraph)

    for subsection in section_data.get("subsections", []):
        document.add_paragraph(subsection["heading"], style="Heading 2")
        for paragraph in subsection.get("paragraphs", []):
            add_body_paragraph(document, paragraph)


def build_document() -> Path:
    document = Document()
    set_page_layout(document)
    configure_styles(document)
    add_footer(document)
    add_title_block(document)

    for section in SECTIONS:
        add_section(document, section)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
